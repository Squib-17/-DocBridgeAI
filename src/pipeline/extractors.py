"""
DocBridgeAI — Content Extractors

One extractor class per file type. All return ExtractedContent so the
rest of the pipeline never needs to know where the content came from.

Extractors:
  MarkdownExtractor    — .md / .txt passthrough
  PDFTextExtractor     — readable PDFs via PyMuPDF
  ScannedPDFExtractor  — image-based PDFs via pytesseract OCR
  DocxExtractor        — Word documents via python-docx
  CSVExtractor         — CSV files via pandas
  XLSXExtractor        — Excel files via pandas + openpyxl

Factory:
  get_extractor(detected) → appropriate extractor instance
"""

from __future__ import annotations

import io
import re
from abc import ABC, abstractmethod
from typing import Any

from .models import DetectedFile, ExtractedContent, SourceFile


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def _table_to_markdown(rows: list[list]) -> str:
    """Render a list of rows (header row first) as a GFM markdown table.

    Single-row detections (e.g. bordered notice boxes in PDFs) are returned
    as plain text to avoid a header-only table with no data rows.
    """
    if not rows:
        return ""
    # Normalize: replace newlines with spaces, coerce None to empty string
    normalized = [
        [(str(c).replace("\n", " ").strip() if c is not None else "") for c in row]
        for row in rows
    ]
    if len(normalized) == 1:
        # Likely a bordered box, not a data table — render as plain text
        return " ".join(c for c in normalized[0] if c)
    header = normalized[0]
    col_count = len(header)
    if col_count == 0:
        return ""
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * col_count) + " |",
    ]
    for row in normalized[1:]:
        padded = (row + [""] * col_count)[:col_count]
        lines.append("| " + " | ".join(padded) + " |")
    return "\n".join(lines)


def _bbox_overlaps(
    a: tuple[float, float, float, float],
    b: tuple[float, float, float, float],
) -> bool:
    """Return True if two axis-aligned bounding boxes overlap."""
    return a[0] < b[2] and a[2] > b[0] and a[1] < b[3] and a[3] > b[1]


# ---------------------------------------------------------------------------
# Base class
# ---------------------------------------------------------------------------

class BaseExtractor(ABC):
    @abstractmethod
    def extract(self, source: SourceFile) -> ExtractedContent:
        ...


# ---------------------------------------------------------------------------
# Markdown / plain text
# ---------------------------------------------------------------------------

class MarkdownExtractor(BaseExtractor):
    """Passthrough extractor for .md and .txt files. Confidence is always 1.0."""

    def extract(self, source: SourceFile) -> ExtractedContent:
        text = source.raw_bytes.decode("utf-8", errors="replace")
        return ExtractedContent(
            source=source,
            file_type="markdown",
            raw_text=text,
            extraction_method="passthrough",
            confidence_hint=1.0,
        )


# ---------------------------------------------------------------------------
# Readable PDF
# ---------------------------------------------------------------------------

class PDFTextExtractor(BaseExtractor):
    """
    Extracts text from a machine-readable PDF using PyMuPDF.
    Strips page headers/footers heuristically and computes a confidence
    hint based on the ratio of extracted text to estimated expected text.
    """

    def extract(self, source: SourceFile) -> ExtractedContent:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=io.BytesIO(source.raw_bytes), filetype="pdf")
        pages_content: list[str] = []

        for page in doc:
            pages_content.append(self._extract_page_content(page))

        doc.close()

        cleaned = self._strip_headers_footers(pages_content)

        char_count = len(cleaned.strip())
        page_count = len(pages_content)
        expected_chars = page_count * 400
        confidence = min(1.0, char_count / max(expected_chars, 1))
        confidence = min(confidence, 0.98)

        return ExtractedContent(
            source=source,
            file_type="pdf_readable",
            raw_text=cleaned,
            page_count=page_count,
            extraction_method="pymupdf",
            confidence_hint=confidence,
        )

    def _extract_page_content(self, page) -> str:
        """
        Return text and tables from one PDF page, interleaved by vertical position.
        Text blocks that overlap table bounding boxes are skipped — their content
        is captured by the table renderer instead.
        """
        try:
            pdf_tables = page.find_tables().tables
        except Exception:
            pdf_tables = []

        table_bboxes = [t.bbox for t in pdf_tables]
        items: list[tuple[float, str]] = []

        for block in page.get_text("blocks"):
            x0, y0, x1, y1, text, _block_no, block_type = block
            if block_type != 0:  # skip image blocks
                continue
            text = text.strip()
            if not text:
                continue
            if any(_bbox_overlaps((x0, y0, x1, y1), tb) for tb in table_bboxes):
                continue
            items.append((y0, text))

        for table in pdf_tables:
            rows = table.extract()
            if rows:
                md = _table_to_markdown(rows)
                if md:
                    items.append((table.bbox[1], md))

        items.sort(key=lambda x: x[0])
        return "\n\n".join(content for _, content in items)

    def _strip_headers_footers(self, pages_text: list[str]) -> str:
        """
        Remove lines that appear identically (or near-identically) on
        more than half the pages — those are almost certainly headers/footers.
        Also remove standalone page number lines.
        """
        if len(pages_text) < 2:
            return pages_text[0] if pages_text else ""

        # Count line frequencies across pages
        from collections import Counter
        line_counts: Counter = Counter()
        for page in pages_text:
            for line in page.splitlines():
                stripped = line.strip()
                if stripped:
                    line_counts[stripped] += 1

        # A line is a header/footer if it appears on MORE THAN half the pages.
        # Using strict > (not >=) prevents over-removal on 2-page docs where
        # threshold=1.0 would otherwise match every line that appears even once.
        threshold = len(pages_text) / 2
        repeated_lines = {line for line, count in line_counts.items() if count > threshold}

        # Page number pattern: a line that is just a number, optionally with "Page X of Y"
        page_number_re = re.compile(r"^\s*(page\s+\d+(\s+of\s+\d+)?|\d+)\s*$", re.IGNORECASE)

        cleaned_pages = []
        for page in pages_text:
            lines = []
            for line in page.splitlines():
                stripped = line.strip()
                if stripped in repeated_lines:
                    continue
                if page_number_re.match(stripped):
                    continue
                lines.append(line)
            cleaned_pages.append("\n".join(lines))

        return "\n\n".join(cleaned_pages)


# ---------------------------------------------------------------------------
# Scanned PDF (OCR)
# ---------------------------------------------------------------------------

class ScannedPDFExtractor(BaseExtractor):
    """
    Converts each PDF page to an image and runs Tesseract OCR.
    Confidence hint is derived from Tesseract's per-word confidence scores.
    """

    def extract(self, source: SourceFile) -> ExtractedContent:
        import fitz  # PyMuPDF — for page-to-image conversion
        import pytesseract
        from PIL import Image

        doc = fitz.open(stream=io.BytesIO(source.raw_bytes), filetype="pdf")
        page_texts: list[str] = []
        all_confidences: list[float] = []

        for page in doc:
            # Render at 300 DPI for good OCR accuracy
            mat = fitz.Matrix(300 / 72, 300 / 72)
            pix = page.get_pixmap(matrix=mat)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Get text with confidence data
            data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
            words = []
            for i, word in enumerate(data["text"]):
                if word.strip():
                    words.append(word)
                    conf = data["conf"][i]
                    if isinstance(conf, (int, float)) and conf >= 0:
                        all_confidences.append(float(conf) / 100.0)

            page_texts.append(" ".join(words))

        doc.close()

        raw_text = "\n\n".join(page_texts)
        avg_confidence = (
            sum(all_confidences) / len(all_confidences) if all_confidences else 0.0
        )

        return ExtractedContent(
            source=source,
            file_type="pdf_scanned",
            raw_text=raw_text,
            page_count=len(page_texts),
            extraction_method="ocr_tesseract",
            confidence_hint=avg_confidence,
        )


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

class DocxExtractor(BaseExtractor):
    """
    Extracts text and tables from a Word document using python-docx.
    Paragraphs and tables are interleaved in document order so that
    table content is not silently dropped from the output.
    DOCX extraction is highly reliable; confidence is set to 1.0.
    """

    def extract(self, source: SourceFile) -> ExtractedContent:
        from docx import Document

        doc = Document(io.BytesIO(source.raw_bytes))

        # Build element-to-object maps so we can iterate body elements in
        # document order while retaining python-docx objects (needed for
        # para.style.name access).
        para_map = {id(p._p): p for p in doc.paragraphs}
        table_map = {id(t._tbl): t for t in doc.tables}

        parts: list[str] = []
        tables: list[dict[str, Any]] = []

        for block in doc.element.body:
            block_id = id(block)

            if block_id in para_map:
                para = para_map[block_id]
                text = para.text.strip()
                if not text:
                    continue
                if para.style.name.startswith("Heading"):
                    try:
                        level = int(para.style.name.split(" ")[-1])
                    except ValueError:
                        level = 2
                    text = "#" * level + " " + text
                parts.append(text)

            elif block_id in table_map:
                table = table_map[block_id]
                rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                if rows:
                    tables.append({"rows": rows, "headers": rows[0]})
                    md = _table_to_markdown(rows)
                    if md:
                        parts.append(md)

        return ExtractedContent(
            source=source,
            file_type="docx",
            raw_text="\n\n".join(parts),
            tables=tables,
            extraction_method="python-docx",
            confidence_hint=1.0,
        )


# ---------------------------------------------------------------------------
# CSV
# ---------------------------------------------------------------------------

class CSVExtractor(BaseExtractor):
    """
    Parses a CSV file using pandas. Returns column names and rows as
    structured data for tabular mode processing.
    """

    def extract(self, source: SourceFile) -> ExtractedContent:
        import pandas as pd

        df = pd.read_csv(io.BytesIO(source.raw_bytes), dtype=str, keep_default_na=False)
        df = df.fillna("")

        column_names = list(df.columns)
        rows = df.to_dict(orient="records")

        # raw_text is a readable preview (first 5 rows)
        preview_df = df.head(5)
        raw_text = preview_df.to_string(index=False)

        return ExtractedContent(
            source=source,
            file_type="csv",
            raw_text=raw_text,
            extraction_method="pandas_csv",
            confidence_hint=1.0,
            column_names=column_names,
            rows=rows,
        )


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------

class XLSXExtractor(BaseExtractor):
    """
    Parses an Excel file using pandas + openpyxl. Reads the first sheet.
    """

    def extract(self, source: SourceFile) -> ExtractedContent:
        import pandas as pd

        df = pd.read_excel(
            io.BytesIO(source.raw_bytes),
            dtype=str,
            keep_default_na=False,
            engine="openpyxl",
        )
        df = df.fillna("")

        column_names = list(df.columns)
        rows = df.to_dict(orient="records")

        preview_df = df.head(5)
        raw_text = preview_df.to_string(index=False)

        return ExtractedContent(
            source=source,
            file_type="xlsx",
            raw_text=raw_text,
            extraction_method="pandas_openpyxl",
            confidence_hint=1.0,
            column_names=column_names,
            rows=rows,
        )


# ---------------------------------------------------------------------------
# Factory
# ---------------------------------------------------------------------------

def get_extractor(detected: DetectedFile) -> BaseExtractor:
    """
    Return the appropriate extractor for a detected file type.
    Raises ValueError for unsupported types — callers should check
    detected.mode != 'unsupported' before calling this.
    """
    mapping: dict[str, type[BaseExtractor]] = {
        "markdown": MarkdownExtractor,
        "pdf_readable": PDFTextExtractor,
        "pdf_scanned": ScannedPDFExtractor,
        "docx": DocxExtractor,
        "csv": CSVExtractor,
        "xlsx": XLSXExtractor,
    }
    extractor_class = mapping.get(detected.file_type)
    if extractor_class is None:
        raise ValueError(
            f"No extractor available for file type: {detected.file_type!r}. "
            "Check that detected.mode != 'unsupported' before calling get_extractor()."
        )
    return extractor_class()
